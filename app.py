import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document

from utils.skills import extract_skills
from utils.pdf_reader import extract_text
from utils.matcher import calculate_similarity
from utils.ats import ats_score
from utils.ai_feedback import generate_feedback

# NEW IMPORTS
from utils.auth import auth_screen, check_login
from utils.database import init_db, insert_result

# --------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------

st.set_page_config(
    page_title="AI TalentMatch Pro",
    page_icon="🚀",
    layout="wide"
)
if not check_login():
    auth_screen()
    st.stop()

# --------------------------------------------------
# PREMIUM UI
# --------------------------------------------------

st.markdown("""
<style>
body {
    background-color: #0e1117;
    color: white;
}
.stButton>button {
    background-color: #4CAF50;
    color: white;
    border-radius: 10px;
}
</style>
""", unsafe_allow_html=True)

# --------------------------------------------------
# LOGIN SYSTEM
# --------------------------------------------------

if not check_login():
    st.warning("🔐 Please login to access the system")
    st.stop()
# -------------------------
# LOGOUT BUTTON
# -------------------------

col1, col2 = st.columns([8,1])

with col2:
    if st.button("🚪 Logout"):
        st.session_state.clear()
        st.rerun()

# --------------------------------------------------
# INIT DATABASE
# --------------------------------------------------

init_db()

# --------------------------------------------------
# SIDEBAR
# --------------------------------------------------

st.sidebar.title("🚀 AI TalentMatch Pro")

st.sidebar.success(f"Logged in as: {st.session_state.get('user')}")

st.sidebar.info("""
### Features

✅ ATS Score  
✅ Resume Ranking  
✅ Skill Gap Analysis  
✅ AI Feedback  
✅ Dashboard Analytics  
✅ Report Download  
""")

# --------------------------------------------------
# TITLE
# --------------------------------------------------

st.title("🚀 AI TalentMatch Pro")

st.markdown("""
### Intelligent Resume Screening & ATS Analysis System

Upload resumes and compare them with a Job Description using AI-powered matching, ATS scoring, and skill gap analysis.
""")

# --------------------------------------------------
# INPUT SECTION
# --------------------------------------------------

job_description = st.text_area("📋 Enter Job Description")

uploaded_files = st.file_uploader(
    "📄 Upload Resume PDFs",
    type=["pdf"],
    accept_multiple_files=True
)

# --------------------------------------------------
# MAIN BUTTON
# --------------------------------------------------

if st.button("🔍 Screen Resumes"):

    if not job_description:
        st.warning("Please enter a Job Description.")

    elif not uploaded_files:
        st.warning("Please upload resumes.")

    else:

        results = []
        jd_skills = extract_skills(job_description)

        for resume in uploaded_files:

            resume_text = extract_text(resume)
            resume_skills = extract_skills(resume_text)

            score = calculate_similarity(resume_text, job_description)

            ats = ats_score(
                resume_text,
                resume_skills,
                jd_skills
            )

            # SAFE AI FEEDBACK (no crash)
            try:
                with st.spinner("🤖 AI analyzing..."):
                    feedback = generate_feedback(resume_text, job_description)
            except:
                feedback = "AI feedback not available (API issue)"

            missing_skills = list(
                set(jd_skills) - set(resume_skills)
            )

            if score >= 80:
                recommendation = "Strong Match"
            elif score >= 60:
                recommendation = "Recommended"
            else:
                recommendation = "Not Recommended"

            # SAVE TO DATABASE
            insert_result(
                resume.name,
                score,
                ats,
                recommendation
            )

            results.append({
                "Resume": resume.name,
                "Match Score (%)": round(score, 2),
                "ATS Score": ats,
                "Skills Found": ", ".join(resume_skills),
                "Missing Skills": ", ".join(missing_skills),
                "Recommendation": recommendation
            })

            # ------------------------------------------
            # RESUME DISPLAY
            # ------------------------------------------

            st.markdown("---")
            st.subheader(f"📄 {resume.name}")

            st.write(f"### Match Score: {round(score,2)}%")
            st.progress(min(int(score), 100))

            col1, col2 = st.columns(2)

            with col1:
                st.metric("ATS Score", f"{ats}/100")

            with col2:
                st.metric("Match Score", f"{round(score,2)}%")

            if score >= 80:
                st.success("🟢 Strong Match")
            elif score >= 60:
                st.warning("🟡 Recommended")
            else:
                st.error("🔴 Not Recommended")

            st.subheader("✅ Skills Found")
            st.write(", ".join(resume_skills) if resume_skills else "No skills detected.")

            st.subheader("❌ Missing Skills")
            st.write(", ".join(missing_skills) if missing_skills else "No Missing Skills")

            st.subheader("🤖 AI Feedback")
            st.write(feedback)

        # ------------------------------------------
        # DATAFRAME
        # ------------------------------------------

        df = pd.DataFrame(results)
        df = df.sort_values(by="Match Score (%)", ascending=False)

        st.markdown("---")
        st.header("🏆 Candidate Ranking")
        st.dataframe(df, use_container_width=True)

        # ------------------------------------------
        # DASHBOARD
        # ------------------------------------------

        st.header("📊 Dashboard Summary")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Total Resumes", len(uploaded_files))

        with col2:
            st.metric("Best Match", f"{df['Match Score (%)'].max()}%")

        with col3:
            st.metric("Average Score", f"{round(df['Match Score (%)'].mean(),2)}%")

        # ------------------------------------------
        # TOP CANDIDATE
        # ------------------------------------------

        best_candidate = df.iloc[0]

        st.success(
            f"🏆 Top Candidate: {best_candidate['Resume']} "
            f"({best_candidate['Match Score (%)']}%)"
        )

        # ------------------------------------------
        # CHART
        # ------------------------------------------

        st.header("📈 Resume Match Score Comparison")

        fig, ax = plt.subplots(figsize=(6, 2))
        ax.barh(df["Resume"], df["Match Score (%)"], height=0.25)
        ax.set_xlim(0, 100)
        ax.set_xlabel("Match Score (%)")
        ax.set_title("Candidate Ranking")
        plt.tight_layout()

        st.pyplot(fig)

        # ------------------------------------------
        # WORD REPORT
        # ------------------------------------------

        document = Document()

        document.add_heading("AI TalentMatch Pro Report", 1)
        document.add_paragraph("Detailed Resume Screening Report")

        document.add_paragraph(f"Total Resumes: {len(uploaded_files)}")
        document.add_paragraph(f"Top Candidate: {best_candidate['Resume']}")
        document.add_paragraph(f"Best Score: {best_candidate['Match Score (%)']}%")

        document.add_heading("Candidate Analysis", 2)

        rank = 1
        for _, row in df.iterrows():

            document.add_heading(f"Rank #{rank}", 3)
            document.add_paragraph(f"Resume: {row['Resume']}")
            document.add_paragraph(f"Match Score: {row['Match Score (%)']}%")
            document.add_paragraph(f"ATS Score: {row['ATS Score']}/100")
            document.add_paragraph(f"Skills: {row['Skills Found']}")
            document.add_paragraph(f"Missing Skills: {row['Missing Skills']}")
            document.add_paragraph(f"Recommendation: {row['Recommendation']}")
            document.add_paragraph("-" * 40)

            rank += 1

        report_file = "Resume_Report.docx"
        document.save(report_file)

        with open(report_file, "rb") as file:
            st.download_button(
                "📄 Download Report",
                file,
                file_name=report_file
            )